"""
Tool registry for the autonomous agent (brain.py).

Each tool = a Gemini FunctionDeclaration (schema the model sees) + a Python impl.
Adding a capability = add one declaration + one function + one dispatch entry.

This is where the agent's real power lives: run_shell can install and run
anything on the machine, so the approval gate (REQUIRE_SHELL_APPROVAL) matters.
"""
import json
import os
import re
import subprocess
import urllib.request
from pathlib import Path

from google.genai import types

_S, _T = types.Schema, types.Type

# Shell approval policy: "risky" (prompt only for installs/deletes/system
# changes), "all" (prompt every command), "never" (fully autonomous).
SHELL_APPROVAL = os.getenv("SHELL_APPROVAL", "risky").lower()
SHELL_TIMEOUT = int(os.getenv("SHELL_TIMEOUT", "300"))
TAVILY_KEY = os.getenv("TAVILY_API_KEY", "")
_MAX_OUT = 8000

# Commands considered risky -> prompt under the "risky" policy. Tuned for Windows
# PowerShell + common cross-platform tools.
_RISKY = re.compile(r"""(?ix)
      \b(install|uninstall|upgrade)\b
    | \b(winget|choco|scoop|pipx|Install-Module)\b
    | \bpip\s+install\b | \bnpm\s+(i|install)\b | \bdotnet\s+tool\b
    | \b(rm|del|erase|rmdir|rd|Remove-Item|Clear-Content)\b
    | \b(format|mkfs|diskpart|bcdedit|fdisk)\b
    | \b(reg|Set-ItemProperty|New-ItemProperty|Remove-ItemProperty)\b
    | \b(netsh|schtasks|setx|New-Service|Set-Service)\b
    | \b(shutdown|restart|Stop-Computer|Restart-Computer|taskkill|Stop-Process)\b
    | \b(Set-ExecutionPolicy|runas|sudo)\b
    | (Invoke-Expression|\|\s*iex\b)
    | git\s+(reset\s+--hard|clean\b|push\s+--force)
""")


def _is_risky(cmd: str) -> bool:
    return bool(_RISKY.search(cmd or ""))


# --------------------------------------------------------------- declarations
DECLARATIONS = [
    types.FunctionDeclaration(
        name="run_shell",
        description=(
            "Run a command on the user's Windows PC and return its combined "
            "stdout/stderr. Use this to install software (winget/pip/npm), run "
            "scripts, manage files, use git, query the system, etc. Commands are "
            "non-interactive — never launch a command that waits for input."),
        parameters=_S(type=_T.OBJECT, properties={
            "command": _S(type=_T.STRING, description="The exact command to run."),
            "shell": _S(type=_T.STRING, enum=["powershell", "bash", "cmd"],
                        description="Interpreter. Default powershell."),
        }, required=["command"]),
    ),
    types.FunctionDeclaration(
        name="read_file",
        description="Read a PLAIN-TEXT file (code, .txt, .md, .json, etc.). For "
                    "PDF/Word documents use read_document instead.",
        parameters=_S(type=_T.OBJECT, properties={
            "path": _S(type=_T.STRING)}, required=["path"]),
    ),
    types.FunctionDeclaration(
        name="read_document",
        description="Extract the text from a document: PDF (.pdf), Word (.docx), "
                    "or plain text. Use this for CVs, resumes, reports — any "
                    "non-plain-text document. read_file cannot read PDF/Word.",
        parameters=_S(type=_T.OBJECT, properties={
            "path": _S(type=_T.STRING)}, required=["path"]),
    ),
    types.FunctionDeclaration(
        name="write_file",
        description="Create or overwrite a text file with the given content.",
        parameters=_S(type=_T.OBJECT, properties={
            "path": _S(type=_T.STRING),
            "content": _S(type=_T.STRING)}, required=["path", "content"]),
    ),
    types.FunctionDeclaration(
        name="list_dir",
        description="List the entries of a directory (default: current dir).",
        parameters=_S(type=_T.OBJECT, properties={
            "path": _S(type=_T.STRING)}),
    ),
    types.FunctionDeclaration(
        name="web_search",
        description=(
            "Search the web (Tavily) and return a short answer plus top results "
            "with titles, URLs, and snippets. Use to look things up online."),
        parameters=_S(type=_T.OBJECT, properties={
            "query": _S(type=_T.STRING),
            "max_results": _S(type=_T.INTEGER, description="Default 5.")},
            required=["query"]),
    ),
    types.FunctionDeclaration(
        name="fetch_url",
        description="Fetch a web page and return its text content (HTML stripped).",
        parameters=_S(type=_T.OBJECT, properties={
            "url": _S(type=_T.STRING)}, required=["url"]),
    ),
    types.FunctionDeclaration(
        name="gmail_recent",
        description=(
            "List recent Gmail messages (read-only) matching a Gmail search "
            "query; returns sender, subject, date, and snippet. Default query "
            "'newer_than:1d'. Use 'is:unread' for unread, 'is:important', etc."),
        parameters=_S(type=_T.OBJECT, properties={
            "query": _S(type=_T.STRING, description="Gmail search query."),
            "max_results": _S(type=_T.INTEGER, description="Default 10.")}),
    ),
    types.FunctionDeclaration(
        name="calendar_events",
        description=(
            "List Google Calendar events (read-only). Default: today. Set 'days' "
            "to widen the window (e.g. 7 for the week)."),
        parameters=_S(type=_T.OBJECT, properties={
            "days": _S(type=_T.INTEGER, description="Days from today. Default 1.")}),
    ),
    types.FunctionDeclaration(
        name="use_computer_visually",
        description=(
            "Perform a task by visually controlling the screen — screenshots plus "
            "real mouse and keyboard: open desktop apps, click buttons, fill "
            "forms, sign in to websites, read on-screen content. Use this ONLY "
            "for GUI tasks that run_shell and the web tools cannot do (e.g. "
            "clicking through an app, a site needing visual interaction). Give a "
            "clear, self-contained task; it returns what was observed/done."),
        parameters=_S(type=_T.OBJECT, properties={
            "task": _S(type=_T.STRING)}, required=["task"]),
    ),
]


# ----------------------------------------------------------------- impls
def _run_shell(args, approve, on_text):
    cmd = args.get("command", "")
    shell = args.get("shell", "powershell")
    need = SHELL_APPROVAL == "all" or (SHELL_APPROVAL == "risky" and _is_risky(cmd))
    if need and not approve(f"[{shell}] {cmd}"):
        return "[user denied this command]"
    tag = "  ⚠️risky" if _is_risky(cmd) else ""
    on_text(f"$ ({shell}){tag} {cmd}")
    try:
        if shell == "powershell":
            argv = ["powershell", "-NoProfile", "-NonInteractive", "-Command", cmd]
        elif shell == "bash":
            argv = ["bash", "-lc", cmd]
        else:
            argv = ["cmd", "/c", cmd]
        out = subprocess.run(argv, capture_output=True, text=True, timeout=SHELL_TIMEOUT)
        body = (out.stdout or "") + (out.stderr or "")
        return body[:_MAX_OUT] or f"[exit {out.returncode}, no output]"
    except subprocess.TimeoutExpired:
        return f"[timed out after {SHELL_TIMEOUT}s]"
    except Exception as e:
        return f"[error: {e}]"


_DOC_EXTS = {".pdf", ".docx", ".doc", ".pptx", ".xlsx", ".odt"}


def _read_file(args, *_):
    p = Path(args["path"])
    if p.suffix.lower() in _DOC_EXTS:
        return (f"[{p.suffix} is a document, not plain text — call read_document "
                f"on this path instead.]")
    try:
        return p.read_text(encoding="utf-8", errors="replace")[:12000]
    except Exception as e:
        return f"[error: {e}]"


def _read_document(args, *_):
    p = Path(args["path"])
    ext = p.suffix.lower()
    try:
        if not p.exists():
            return f"[no such file: {p}]"
        if ext == ".pdf":
            from pypdf import PdfReader
            text = "\n".join((pg.extract_text() or "") for pg in PdfReader(str(p)).pages)
        elif ext == ".docx":
            from docx import Document
            doc = Document(str(p))
            parts = [para.text for para in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            text = "\n".join(parts)
        else:
            text = p.read_text(encoding="utf-8", errors="replace")
        text = text.strip()
        return text[:20000] or "[document has no extractable text]"
    except Exception as e:
        return f"[document error: {e}]"


def _write_file(args, *_):
    try:
        p = Path(args["path"])
        if p.parent and not p.parent.exists():
            p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(args.get("content", ""), encoding="utf-8")
        return f"wrote {p}"
    except Exception as e:
        return f"[error: {e}]"


def _list_dir(args, *_):
    try:
        p = Path(args.get("path", "."))
        return "\n".join(sorted(
            x.name + ("/" if x.is_dir() else "") for x in p.iterdir())) or "[empty]"
    except Exception as e:
        return f"[error: {e}]"


def _web_search(args, *_):
    if not TAVILY_KEY:
        return "[web_search unavailable: set TAVILY_API_KEY in .env]"
    payload = json.dumps({
        "api_key": TAVILY_KEY,
        "query": args.get("query", ""),
        "max_results": int(args.get("max_results", 5)),
        "include_answer": True,
    }).encode()
    req = urllib.request.Request(
        "https://api.tavily.com/search", data=payload,
        headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=30) as r:
            data = json.loads(r.read())
        lines = []
        if data.get("answer"):
            lines.append("Answer: " + data["answer"])
        for it in data.get("results", []):
            lines.append(f"- {it.get('title', '')} ({it.get('url', '')})\n"
                         f"  {(it.get('content') or '')[:300]}")
        return "\n".join(lines) or "[no results]"
    except Exception as e:
        return f"[search error: {e}]"


def _fetch_url(args, *_):
    url = args.get("url", "")
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as r:
            raw = r.read(2_000_000).decode("utf-8", "replace")
        text = re.sub(r"(?is)<(script|style)\b.*?</\1>", " ", raw)
        text = re.sub(r"(?s)<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text[:_MAX_OUT] or "[empty page]"
    except Exception as e:
        return f"[fetch error: {e}]"


def _use_computer(args, approve, on_text):
    """Delegate to the Gemini computer-use loop (screenshot -> click). Imported
    lazily so pure shell/web sessions don't load pyautogui."""
    task = args.get("task", "")
    try:
        import gemini_agent
        session = gemini_agent.new_session()
        gemini_agent.add_user_text(
            session, task + "\nWhen finished, clearly state the result in text.")
        captured = []

        def tee(t):
            captured.append(t)
            on_text("[gui] " + t)

        gemini_agent.run_turn(session, tee, approve)
        return "\n".join(captured)[-4000:] or "[gui task finished]"
    except Exception as e:
        return f"[gui error: {e}]"


def _gmail_recent(args, approve, on_text):
    on_text("📧 reading recent Gmail…")
    import google_tools
    return google_tools.gmail_recent(args)


def _calendar_events(args, approve, on_text):
    on_text("📅 reading calendar…")
    import google_tools
    return google_tools.calendar_events(args)


_IMPL = {
    "run_shell": _run_shell,
    "read_file": _read_file,
    "read_document": _read_document,
    "write_file": _write_file,
    "list_dir": _list_dir,
    "web_search": _web_search,
    "fetch_url": _fetch_url,
    "gmail_recent": _gmail_recent,
    "calendar_events": _calendar_events,
    "use_computer_visually": _use_computer,
}


def dispatch(name, args, approve, on_text):
    fn = _IMPL.get(name)
    if fn is None:
        return f"[unknown tool: {name}]"
    return fn(args, approve, on_text)
