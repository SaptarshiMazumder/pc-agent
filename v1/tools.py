"""
Tool registry for the autonomous agent (brain.py).

Each tool = a Gemini FunctionDeclaration (schema the model sees) + a Python impl.
Adding a capability = add one declaration + one function + one dispatch entry.

This is where the agent's real power lives: run_shell can install and run
anything on the machine, so the approval gate (REQUIRE_SHELL_APPROVAL) matters.
"""
import html as _html
import json
import os
import re
import subprocess
import threading
import time
import urllib.error
import urllib.request
from pathlib import Path

from google.genai import types

_S, _T = types.Schema, types.Type

# Shell approval policy: "risky" (prompt only for installs/deletes/system
# changes), "all" (prompt every command), "never" (fully autonomous).
SHELL_APPROVAL = os.getenv("SHELL_APPROVAL", "risky").lower()
SHELL_TIMEOUT = int(os.getenv("SHELL_TIMEOUT", "300"))
TAVILY_KEY = os.getenv("TAVILY_API_KEY", "")
_MAX_OUT = 8000

# Commands considered risky -> prompt under the "risky" policy. Tuned for Windows
# PowerShell + common cross-platform tools.
_RISKY = re.compile(r"""(?ix)
      \b(install|uninstall|upgrade)\b
    | \b(winget|choco|scoop|pipx|Install-Module)\b
    | \bpip\s+install\b | \bnpm\s+(i|install)\b | \bdotnet\s+tool\b
    | \b(rm|del|erase|rmdir|rd|Remove-Item|Clear-Content)\b
    | \b(format|mkfs|diskpart|bcdedit|fdisk)\b
    | \b(reg|Set-ItemProperty|New-ItemProperty|Remove-ItemProperty)\b
    | \b(netsh|schtasks|setx|New-Service|Set-Service)\b
    | \b(shutdown|restart|Stop-Computer|Restart-Computer|taskkill|Stop-Process)\b
    | \b(Set-ExecutionPolicy|runas|sudo)\b
    | (Invoke-Expression|\|\s*iex\b)
    | git\s+(reset\s+--hard|clean\b|push\s+--force)
""")


def _is_risky(cmd: str) -> bool:
    return bool(_RISKY.search(cmd or ""))


# --------------------------------------------------------------- declarations
DECLARATIONS = [
    types.FunctionDeclaration(
        name="run_shell",
        description=(
            "Run a command on the user's Windows PC and return its combined "
            "stdout/stderr. Use this to install software (winget/pip/npm), run "
            "scripts, manage files, use git, query the system, etc. Commands are "
            "non-interactive — never launch a command that waits for input."),
        parameters=_S(type=_T.OBJECT, properties={
            "command": _S(type=_T.STRING, description="The exact command to run."),
            "shell": _S(type=_T.STRING, enum=["powershell", "bash", "cmd"],
                        description="Interpreter. Default powershell."),
        }, required=["command"]),
    ),
    types.FunctionDeclaration(
        name="read_file",
        description="Read a PLAIN-TEXT file (code, .txt, .md, .json, etc.). For "
                    "PDF/Word documents use read_document instead.",
        parameters=_S(type=_T.OBJECT, properties={
            "path": _S(type=_T.STRING)}, required=["path"]),
    ),
    types.FunctionDeclaration(
        name="read_document",
        description="Extract the text from a document: PDF (.pdf), Word (.docx), "
                    "or plain text. Use this for CVs, resumes, reports — any "
                    "non-plain-text document. read_file cannot read PDF/Word.",
        parameters=_S(type=_T.OBJECT, properties={
            "path": _S(type=_T.STRING)}, required=["path"]),
    ),
    types.FunctionDeclaration(
        name="write_file",
        description="Create or overwrite a text file with the given content.",
        parameters=_S(type=_T.OBJECT, properties={
            "path": _S(type=_T.STRING),
            "content": _S(type=_T.STRING)}, required=["path", "content"]),
    ),
    types.FunctionDeclaration(
        name="list_dir",
        description="List the entries of a directory (default: current dir).",
        parameters=_S(type=_T.OBJECT, properties={
            "path": _S(type=_T.STRING)}),
    ),
    types.FunctionDeclaration(
        name="find_file",
        description=(
            "Find files by name (case-insensitive substring) across the user's "
            "common folders — Desktop, OneDrive Desktop, Documents, Downloads. "
            "Fast and bounded. ALWAYS use this to locate a file; do NOT run "
            "recursive directory scans (Get-ChildItem -Recurse over the home "
            "folder), which can hang on OneDrive or huge trees."),
        parameters=_S(type=_T.OBJECT, properties={
            "name": _S(type=_T.STRING, description="Filename or substring to find.")},
            required=["name"]),
    ),
    types.FunctionDeclaration(
        name="web_search",
        description=(
            "Search the web (Tavily) and return a short answer plus top results "
            "with titles, URLs, and snippets. Use to look things up online."),
        parameters=_S(type=_T.OBJECT, properties={
            "query": _S(type=_T.STRING),
            "max_results": _S(type=_T.INTEGER, description="Default 5.")},
            required=["query"]),
    ),
    types.FunctionDeclaration(
        name="fetch_url",
        description="Fetch a web page and return its text content (HTML stripped).",
        parameters=_S(type=_T.OBJECT, properties={
            "url": _S(type=_T.STRING)}, required=["url"]),
    ),
    types.FunctionDeclaration(
        name="gmail_recent",
        description=(
            "List recent Gmail messages (read-only) matching a Gmail search "
            "query; returns sender, subject, date, and snippet. Default query "
            "'newer_than:1d'. Use 'is:unread' for unread, 'is:important', etc."),
        parameters=_S(type=_T.OBJECT, properties={
            "query": _S(type=_T.STRING, description="Gmail search query."),
            "max_results": _S(type=_T.INTEGER, description="Default 10.")}),
    ),
    types.FunctionDeclaration(
        name="calendar_events",
        description=(
            "List Google Calendar events (read-only). Default: today. Set 'days' "
            "to widen the window (e.g. 7 for the week)."),
        parameters=_S(type=_T.OBJECT, properties={
            "days": _S(type=_T.INTEGER, description="Days from today. Default 1.")}),
    ),
    # --- web browser (agent-browser: CDP + accessibility-tree @refs) ---
    # Core loop: web_open -> web_snapshot (read @eN refs) -> act on a ref
    # (web_click/web_fill/web_press) -> re-snapshot. The browser stays warm
    # across these calls. Use these for ALL web work instead of the visual tool.
    types.FunctionDeclaration(
        name="web_open",
        description=(
            "Open a URL and return a CONTENT snapshot of the page: the readable TEXT "
            "(headings, paragraphs, list items — e.g. a full job description) PLUS "
            "interactive @eN refs and link URLs. START any web task here. To FIND or "
            "LIST many items, open the site's SEARCH-RESULTS URL (with query params) "
            "and read the whole list from this one snapshot — do NOT open results one "
            "link at a time."),
        parameters=_S(type=_T.OBJECT, properties={
            "url": _S(type=_T.STRING)}, required=["url"]),
    ),
    types.FunctionDeclaration(
        name="web_snapshot",
        description=(
            "Re-take the CONTENT snapshot of the CURRENT page: readable text + @eN "
            "refs + link URLs. Refs go STALE after the page changes — call this again "
            "after every click/navigation before using a ref. Reading this snapshot "
            "is usually enough to extract the data; you rarely need web_get_text."),
        parameters=_S(type=_T.OBJECT, properties={}),
    ),
    types.FunctionDeclaration(
        name="web_click",
        description="Click an element by its @eN ref (from the latest snapshot). "
                    "Returns a fresh snapshot.",
        parameters=_S(type=_T.OBJECT, properties={
            "ref": _S(type=_T.STRING, description="e.g. @e5")}, required=["ref"]),
    ),
    types.FunctionDeclaration(
        name="web_fill",
        description="Clear a field (by @eN ref) and type text into it.",
        parameters=_S(type=_T.OBJECT, properties={
            "ref": _S(type=_T.STRING), "text": _S(type=_T.STRING)},
            required=["ref", "text"]),
    ),
    types.FunctionDeclaration(
        name="web_press",
        description="Press a key at the current focus (e.g. 'Enter', 'Tab', "
                    "'Control+a'). Returns a fresh snapshot.",
        parameters=_S(type=_T.OBJECT, properties={
            "key": _S(type=_T.STRING)}, required=["key"]),
    ),
    types.FunctionDeclaration(
        name="web_get_text",
        description="Get the visible text of an element by its @eN ref.",
        parameters=_S(type=_T.OBJECT, properties={
            "ref": _S(type=_T.STRING)}, required=["ref"]),
    ),
    types.FunctionDeclaration(
        name="web_scroll",
        description="Scroll the page ('up'/'down'/'left'/'right'), optional pixel "
                    "amount. Returns a fresh snapshot.",
        parameters=_S(type=_T.OBJECT, properties={
            "direction": _S(type=_T.STRING), "pixels": _S(type=_T.INTEGER)},
            required=["direction"]),
    ),
    types.FunctionDeclaration(
        name="web_login",
        description=(
            "When web_open reports '[NOT LOGGED IN]' or a sign-in wall, call this "
            "to let the USER log in: it opens a VISIBLE browser window (on the "
            "persistent profile) at the given login URL and waits for the user to "
            "finish. After it returns, the session is saved — retry web_open and "
            "you'll be authenticated. Pass the site's login URL."),
        parameters=_S(type=_T.OBJECT, properties={
            "url": _S(type=_T.STRING, description="Login page URL, e.g. https://www.linkedin.com/login")}),
    ),
    types.FunctionDeclaration(
        name="web_close",
        description="Close the browser when the web task is done.",
        parameters=_S(type=_T.OBJECT, properties={}),
    ),
    types.FunctionDeclaration(
        name="use_computer_visually",
        description=(
            "LAST RESORT, ~10x slower. Visually control the screen with screenshots "
            "+ real mouse/keyboard. Use ONLY for NATIVE DESKTOP applications (not "
            "websites) — e.g. clicking through a desktop app's UI. For ANYTHING on "
            "the web, use browse_web instead. Give a clear, self-contained task."),
        parameters=_S(type=_T.OBJECT, properties={
            "task": _S(type=_T.STRING)}, required=["task"]),
    ),
]


# ----------------------------------------------------------------- impls
def _run_shell(args, approve, on_text):
    cmd = args.get("command", "")
    shell = args.get("shell", "powershell")
    need = SHELL_APPROVAL == "all" or (SHELL_APPROVAL == "risky" and _is_risky(cmd))
    if need and not approve(f"[{shell}] {cmd}"):
        return "[user denied this command]"
    tag = "  ⚠️risky" if _is_risky(cmd) else ""
    on_text(f"$ ({shell}){tag} {cmd}")
    try:
        if shell == "powershell":
            argv = ["powershell", "-NoProfile", "-NonInteractive", "-Command", cmd]
        elif shell == "bash":
            argv = ["bash", "-lc", cmd]
        else:
            argv = ["cmd", "/c", cmd]
        out = subprocess.run(argv, capture_output=True, text=True, timeout=SHELL_TIMEOUT)
        body = (out.stdout or "") + (out.stderr or "")
        return body[:_MAX_OUT] or f"[exit {out.returncode}, no output]"
    except subprocess.TimeoutExpired:
        return f"[timed out after {SHELL_TIMEOUT}s]"
    except Exception as e:
        return f"[error: {e}]"


_DOC_EXTS = {".pdf", ".docx", ".doc", ".pptx", ".xlsx", ".odt"}


def _read_file(args, *_):
    p = Path(args["path"])
    if p.suffix.lower() in _DOC_EXTS:
        return (f"[{p.suffix} is a document, not plain text — call read_document "
                f"on this path instead.]")
    try:
        return p.read_text(encoding="utf-8", errors="replace")[:12000]
    except Exception as e:
        return f"[error: {e}]"


def _read_document(args, *_):
    p = Path(args["path"])
    ext = p.suffix.lower()
    try:
        if not p.exists():
            return f"[no such file: {p}]"
        if ext == ".pdf":
            from pypdf import PdfReader
            text = "\n".join((pg.extract_text() or "") for pg in PdfReader(str(p)).pages)
        elif ext == ".docx":
            from docx import Document
            doc = Document(str(p))
            parts = [para.text for para in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            text = "\n".join(parts)
        else:
            text = p.read_text(encoding="utf-8", errors="replace")
        text = text.strip()
        return text[:20000] or "[document has no extractable text]"
    except Exception as e:
        return f"[document error: {e}]"


def _write_file(args, *_):
    try:
        p = Path(args["path"])
        if p.parent and not p.parent.exists():
            p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(args.get("content", ""), encoding="utf-8")
        return f"wrote {p}"
    except Exception as e:
        return f"[error: {e}]"


def _list_dir(args, *_):
    try:
        p = Path(args.get("path", "."))
        return "\n".join(sorted(
            x.name + ("/" if x.is_dir() else "") for x in p.iterdir())) or "[empty]"
    except Exception as e:
        return f"[error: {e}]"


_FIND_SKIP = {"node_modules", ".git", ".venv", "venv", "__pycache__", "AppData",
              "$Recycle.Bin", "Windows", "Program Files", "Program Files (x86)"}


def _find_file(args, *_):
    """Bounded file search across common user folders — no full-disk recursion."""
    pat = (args.get("name") or "").lower()
    if not pat:
        return "[provide a name or substring to search for]"
    home = Path.home()
    roots, seen = [], set()
    for p in (home / "Desktop", home / "OneDrive" / "Desktop",
              home / "Documents", home / "OneDrive" / "Documents",
              home / "Downloads"):
        if p.exists() and p not in seen:
            seen.add(p)
            roots.append(p)
    hits, deadline = [], time.time() + 15
    for root in roots:
        base = len(root.parts)
        for dirpath, dirnames, filenames in os.walk(root):
            if time.time() > deadline or len(hits) >= 50:
                break
            if len(Path(dirpath).parts) - base >= 6:
                dirnames[:] = []                      # cap depth
            dirnames[:] = [d for d in dirnames
                           if d not in _FIND_SKIP and not d.startswith(".")]
            for f in filenames:
                if pat in f.lower():
                    hits.append(str(Path(dirpath) / f))
        if len(hits) >= 50 or time.time() > deadline:
            break
    if not hits:
        return f"[no files matching '{pat}' in Desktop/Documents/Downloads]"
    return "\n".join(hits[:50]) + ("\n…(more, narrow the name)" if len(hits) >= 50 else "")


def _web_search(args, *_):
    if not TAVILY_KEY:
        return "[web_search unavailable: set TAVILY_API_KEY in .env]"
    payload = json.dumps({
        "api_key": TAVILY_KEY,
        "query": args.get("query", ""),
        "max_results": int(args.get("max_results", 5)),
        "include_answer": True,
    }).encode()
    req = urllib.request.Request(
        "https://api.tavily.com/search", data=payload,
        headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=30) as r:
            data = json.loads(r.read())
        lines = []
        if data.get("answer"):
            lines.append("Answer: " + data["answer"])
        for it in data.get("results", []):
            lines.append(f"- {it.get('title', '')} ({it.get('url', '')})\n"
                         f"  {(it.get('content') or '')[:300]}")
        return "\n".join(lines) or "[no results]"
    except Exception as e:
        return f"[search error: {e}]"


# A real Chrome UA + browser-like Accept headers. A bare urllib User-Agent
# ("Python-urllib", or even "Mozilla/5.0") gets 403'd by most major sites'
# anti-bot rules; these headers mimic a genuine browser request (this is how
# OpenClaw's web_fetch avoids 403s — see reference web-fetch.ts).
_BROWSER_HEADERS = {
    "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                   "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"),
    "Accept": ("text/html,application/xhtml+xml,application/xml;q=0.9,"
               "image/avif,image/webp,*/*;q=0.8"),
    "Accept-Language": "en-US,en;q=0.9",
}

# Boilerplate containers stripped before extracting readable text — a tiny
# Readability-lite pass so we return the page's CONTENT, not nav/script noise.
_BOILERPLATE = ("script", "style", "noscript", "template", "svg", "head",
                "nav", "header", "footer", "aside", "form")


def _html_to_text(raw):
    text = raw
    for tag in _BOILERPLATE:
        text = re.sub(rf"(?is)<{tag}\b.*?</{tag}>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)        # drop remaining tags
    text = _html.unescape(text)                     # &amp; -> &, &#39; -> '
    text = re.sub(r"[ \t\r\f]+", " ", text)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    return text.strip()


def _fetch_url(args, *_):
    url = args.get("url", "")
    if "://" not in url:
        url = "https://" + url
    try:
        req = urllib.request.Request(url, headers=_BROWSER_HEADERS)
        with urllib.request.urlopen(req, timeout=30) as r:
            raw = r.read(3_000_000).decode("utf-8", "replace")
        text = _html_to_text(raw)
        if not text:
            return "[empty page]"
        if len(text) > _MAX_OUT:
            text = text[:_MAX_OUT] + "\n…[truncated]"
        return text
    except urllib.error.HTTPError as e:             # 403/401/429 -> the site is walling us
        return (f"[fetch error: HTTP {e.code} {e.reason}. This site blocks plain "
                f"HTTP — open it with the web browser (web_open) instead.]")
    except Exception as e:
        return f"[fetch error: {e}]"


def _web_open(args, approve, on_text):
    on_text(f"🌐 open {args.get('url', '')}")
    import agent_browser as ab
    return ab.open_url(args.get("url", ""))


def _web_snapshot(args, approve, on_text):
    import agent_browser as ab
    return ab.snapshot()


def _web_click(args, approve, on_text):
    on_text(f"🌐 click {args.get('ref', '')}")
    import agent_browser as ab
    return ab.click(args.get("ref", ""))


def _web_fill(args, approve, on_text):
    on_text(f"🌐 fill {args.get('ref', '')}")
    import agent_browser as ab
    return ab.fill(args.get("ref", ""), args.get("text", ""))


def _web_press(args, approve, on_text):
    on_text(f"🌐 press {args.get('key', '')}")
    import agent_browser as ab
    return ab.press(args.get("key", ""))


def _web_get_text(args, approve, on_text):
    import agent_browser as ab
    return ab.get_text(args.get("ref", ""))


def _web_scroll(args, approve, on_text):
    import agent_browser as ab
    return ab.scroll(args.get("direction", "down"), args.get("pixels"))


def _web_login(args, approve, on_text):
    import agent_browser as ab
    on_text(ab.login(args.get("url", "")))
    on_text("➡️  Log in to the site in the browser window that just opened.")
    approve("Confirm once you have finished logging in")   # blocks for the user
    ab.close()
    return ("[login window closed — your session is saved to the profile. "
            "Now retry web_open; you should be authenticated.]")


def _web_close(args, approve, on_text):
    import agent_browser as ab
    return ab.close()


def _use_computer(args, approve, on_text):
    """Delegate to the Gemini computer-use loop (screenshot -> click). Imported
    lazily so pure shell/web sessions don't load pyautogui."""
    task = args.get("task", "")
    try:
        import gemini_agent
        session = gemini_agent.new_session()
        gemini_agent.add_user_text(
            session, task + "\nWhen finished, clearly state the result in text.")
        captured = []

        def tee(t):
            captured.append(t)
            on_text("[gui] " + t)

        gemini_agent.run_turn(session, tee, approve)
        return "\n".join(captured)[-4000:] or "[gui task finished]"
    except Exception as e:
        return f"[gui error: {e}]"


def _gmail_recent(args, approve, on_text):
    on_text("📧 reading recent Gmail…")
    import google_tools
    return google_tools.gmail_recent(args)


def _calendar_events(args, approve, on_text):
    on_text("📅 reading calendar…")
    import google_tools
    return google_tools.calendar_events(args)


_IMPL = {
    "run_shell": _run_shell,
    "read_file": _read_file,
    "read_document": _read_document,
    "write_file": _write_file,
    "list_dir": _list_dir,
    "find_file": _find_file,
    "web_search": _web_search,
    "fetch_url": _fetch_url,
    "gmail_recent": _gmail_recent,
    "calendar_events": _calendar_events,
    "web_open": _web_open,
    "web_snapshot": _web_snapshot,
    "web_click": _web_click,
    "web_fill": _web_fill,
    "web_press": _web_press,
    "web_get_text": _web_get_text,
    "web_scroll": _web_scroll,
    "web_login": _web_login,
    "web_close": _web_close,
    "use_computer_visually": _use_computer,
}


# --- watchdog: no single tool call may freeze the agent ---
# Per-tool time budget (seconds). A tool that exceeds it returns a timeout
# message so the model can recover, instead of hanging the whole run. Note: a
# timed-out call's background thread can't be force-killed in Python (Windows has
# no signal.alarm), so it lingers until it finishes — but the AGENT is unblocked.
# run_shell is exempt from this leak: its subprocess is killed by its own timeout.
TOOL_TIMEOUT = int(os.getenv("TOOL_TIMEOUT", "120"))
_TOOL_TIMEOUTS = {
    "run_shell": SHELL_TIMEOUT + 30,                       # subprocess timeout fires first
    "use_computer_visually": int(os.getenv("GUI_TIMEOUT", "600")),
    "web_open": 120,                                       # daemon start + navigation
    "web_login": 600,                                      # user takes time to log in
    "gmail_recent": 300,                                   # allow first-run OAuth consent
    "calendar_events": 300,
}


def _guard(name, fn, args, approve, on_text):
    timeout = _TOOL_TIMEOUTS.get(name, TOOL_TIMEOUT)
    box = {}

    def work():
        try:
            box["r"] = fn(args, approve, on_text)
        except Exception as e:                             # noqa: BLE001
            box["e"] = e

    t = threading.Thread(target=work, daemon=True)
    t.start()
    # Poll in short slices instead of one long join(): on Windows a single long
    # join() blocks Ctrl+C until it returns, so the user can't interrupt a slow
    # tool. Short joins let the interpreter raise KeyboardInterrupt between slices
    # (it propagates out and stops the run; the daemon worker is abandoned).
    deadline = time.monotonic() + timeout
    while t.is_alive() and time.monotonic() < deadline:
        t.join(0.2)
    if t.is_alive():
        return (f"[{name} timed out after {timeout}s and may be stuck — do NOT "
                f"retry it the same way; try a different tool or approach.]")
    if "e" in box:
        return f"[{name} error: {box['e']}]"
    return box.get("r", "")


def dispatch(name, args, approve, on_text):
    fn = _IMPL.get(name)
    if fn is None:
        return f"[unknown tool: {name}]"
    return _guard(name, fn, args, approve, on_text)
