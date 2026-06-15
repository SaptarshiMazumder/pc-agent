"""File tools: read, write, edit.

Mirrors the reference coding tools:
  read  — path, offset (1-indexed), limit; truncation flags
  write — mkdir -p, utf-8 write
  edit  — list of {oldText, newText}; each oldText must match exactly once,
          LF-normalized, non-overlapping; returns a unified diff
"""

from __future__ import annotations

import asyncio
import difflib
import fnmatch
import os
import sys
from pathlib import Path

from . import Tool, ToolResult

MAX_READ_CHARS = 100_000
MAX_LINE_CHARS = 2_000
# Recursive search skips these noise/heavy/system dirs for speed (does NOT
# restrict access — any absolute path is still readable, and you can search any
# root). These are generic folder NAMES, identical on every machine.
FIND_SKIP_DIRS = {
    # caches / dev noise
    "appdata", "node_modules", ".git", ".venv", "venv", "__pycache__",
    ".cache", "$recycle.bin", "system volume information",
    ".vscode", ".cursor", ".idea", "site-packages", ".npm", ".nuget",
    # OS / program dirs (only relevant when scanning a whole drive)
    "windows", "program files", "program files (x86)", "programdata",
    "perflogs", "recovery", "$windows.~bt", "$windows.~ws",
}


def _drive_roots() -> list[Path]:
    """All fixed drive roots (C:\\, D:\\, …) on Windows; '/' elsewhere."""
    if sys.platform == "win32":
        import string

        return [Path(f"{d}:\\") for d in string.ascii_uppercase if os.path.exists(f"{d}:\\")]
    return [Path("/")]
# Real-document extensions are ranked first so noise can't bury an actual file.
_DOC_EXTS = {
    ".docx", ".doc", ".pdf", ".txt", ".md", ".rtf", ".odt",
    ".xlsx", ".xls", ".pptx", ".ppt", ".csv", ".pages",
}
_FIND_HARD_CAP = 5000  # bound walk time; we collect then rank then trim


def _resolve(config, path: str) -> Path:
    p = Path(path)
    if not p.is_absolute():
        p = Path(config.workspace) / p
    return p


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


class ReadTool(Tool):
    name = "read"
    description = (
        "Read a file's contents (text, code, .docx, or .pdf — documents are text-extracted). "
        "Use an absolute path for files outside the workspace. Supports offset (1-indexed line) and limit."
    )
    label = "Read"
    parameters = {
        "type": "object",
        "required": ["path"],
        "properties": {
            "path": {"type": "string", "description": "File path (absolute or workspace-relative)."},
            "offset": {"type": "integer", "minimum": 1, "description": "1-indexed start line."},
            "limit": {"type": "integer", "minimum": 1, "description": "Max lines to read."},
        },
    }

    def __init__(self, config):
        self.config = config

    async def execute(self, tool_call_id, params, abort, on_update=None):
        path = _resolve(self.config, params["path"])
        if not path.is_file():
            return ToolResult.text(f"File not found: {path}", is_error=True)

        # Documents: extract text instead of reading raw bytes.
        suffix = path.suffix.lower()
        if suffix in (".docx", ".pdf"):
            try:
                extractor = _extract_docx if suffix == ".docx" else _extract_pdf
                text = await asyncio.to_thread(extractor, path)
            except Exception as e:
                return ToolResult.text(
                    f"Could not extract {suffix} ({type(e).__name__}: {e})", is_error=True
                )
            if len(text) > MAX_READ_CHARS:
                text = text[:MAX_READ_CHARS] + "\n… [output truncated]"
            return ToolResult.text(text or f"(no extractable text in {path.name})")

        try:
            raw = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return ToolResult.text(f"Cannot read binary file as text: {path}", is_error=True)

        lines = raw.splitlines()
        offset = params.get("offset", 1)
        limit = params.get("limit")
        selected = lines[offset - 1 : offset - 1 + limit] if limit else lines[offset - 1 :]
        truncated_lines = len(selected) < len(lines) - (offset - 1)

        out_lines = []
        for i, line in enumerate(selected, start=offset):
            if len(line) > MAX_LINE_CHARS:
                line = line[:MAX_LINE_CHARS] + "… [line truncated]"
            out_lines.append(f"{i}\t{line}")
        text = "\n".join(out_lines)
        if len(text) > MAX_READ_CHARS:
            text = text[:MAX_READ_CHARS] + "\n… [output truncated]"
            truncated_lines = True
        if truncated_lines:
            text += f"\n(showing {len(selected)} of {len(lines)} lines)"
        return ToolResult.text(text if text else "(empty file)")


class WriteTool(Tool):
    name = "write"
    description = "Create or overwrite a file with the given content."
    label = "Write"
    concurrency = "sequential"
    parameters = {
        "type": "object",
        "required": ["path", "content"],
        "properties": {
            "path": {"type": "string", "description": "File path (absolute or workspace-relative)."},
            "content": {"type": "string", "description": "Full file content."},
        },
    }

    def __init__(self, config):
        self.config = config

    async def execute(self, tool_call_id, params, abort, on_update=None):
        path = _resolve(self.config, params["path"])
        existed = path.exists()
        path.parent.mkdir(parents=True, exist_ok=True)
        content = params["content"]
        await asyncio.to_thread(path.write_text, content, "utf-8")
        verb = "Overwrote" if existed else "Created"
        return ToolResult.text(f"{verb} {path} ({len(content.encode('utf-8'))} bytes)")


def apply_edits(original: str, edits: list[dict]) -> tuple[str, str]:
    """Apply [{oldText, newText}] edits. Each oldText must occur exactly once
    (after LF normalization); spans must not overlap. Returns (new_text, diff).
    Raises ValueError with a helpful message otherwise."""
    had_crlf = "\r\n" in original
    text = original.replace("\r\n", "\n")

    spans: list[tuple[int, int, str]] = []
    for i, edit in enumerate(edits):
        old = edit["oldText"].replace("\r\n", "\n")
        new = edit["newText"].replace("\r\n", "\n")
        if not old:
            raise ValueError(f"Edit {i + 1}: oldText must not be empty.")
        count = text.count(old)
        if count == 0:
            close = difflib.get_close_matches(
                old.strip().splitlines()[0] if old.strip() else "", text.splitlines(), n=1, cutoff=0.5
            )
            hint = f" Closest line in file: {close[0]!r}." if close else ""
            raise ValueError(
                f"Edit {i + 1}: oldText not found in file.{hint} "
                f"File starts with: {text[:300]!r}"
            )
        if count > 1:
            raise ValueError(
                f"Edit {i + 1}: oldText matches {count} times; it must be unique. "
                "Add more surrounding context."
            )
        start = text.index(old)
        spans.append((start, start + len(old), new))

    spans.sort(key=lambda s: s[0])
    for (s1, e1, _), (s2, _e2, _) in zip(spans, spans[1:]):
        if s2 < e1:
            raise ValueError("Edits overlap; make them non-overlapping.")

    result = []
    cursor = 0
    for start, end, new in spans:
        result.append(text[cursor:start])
        result.append(new)
        cursor = end
    result.append(text[cursor:])
    new_text = "".join(result)

    diff = "\n".join(
        difflib.unified_diff(
            text.splitlines(), new_text.splitlines(), fromfile="before", tofile="after", lineterm=""
        )
    )
    if had_crlf:
        new_text = new_text.replace("\n", "\r\n")
    return new_text, diff


class EditTool(Tool):
    name = "edit"
    description = "Make precise text replacements in a file. Each oldText must match exactly once."
    label = "Edit"
    concurrency = "sequential"
    parameters = {
        "type": "object",
        "required": ["path", "edits"],
        "properties": {
            "path": {"type": "string", "description": "File path (absolute or workspace-relative)."},
            "edits": {
                "type": "array",
                "minItems": 1,
                "items": {
                    "type": "object",
                    "required": ["oldText", "newText"],
                    "properties": {
                        "oldText": {"type": "string", "description": "Exact text to replace (unique in file)."},
                        "newText": {"type": "string", "description": "Replacement text."},
                    },
                },
            },
        },
    }

    def __init__(self, config):
        self.config = config

    async def execute(self, tool_call_id, params, abort, on_update=None):
        path = _resolve(self.config, params["path"])
        if not path.is_file():
            return ToolResult.text(f"File not found: {path}", is_error=True)
        original = path.read_text(encoding="utf-8")
        try:
            new_text, diff = apply_edits(original, params["edits"])
        except ValueError as e:
            return ToolResult.text(str(e), is_error=True)
        await asyncio.to_thread(path.write_text, new_text, "utf-8")
        return ToolResult.text(f"Edited {path}:\n{diff}")


class LsTool(Tool):
    name = "ls"
    description = "List the contents of a directory (cross-platform). Defaults to the workspace."
    label = "List"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Directory path (absolute or workspace-relative)."},
        },
    }

    def __init__(self, config):
        self.config = config

    async def execute(self, tool_call_id, params, abort, on_update=None):
        path = _resolve(self.config, params.get("path") or ".")
        if not path.exists():
            return ToolResult.text(f"Directory not found: {path}", is_error=True)
        if not path.is_dir():
            return ToolResult.text(f"Not a directory: {path}", is_error=True)
        try:
            entries = sorted(path.iterdir(), key=lambda p: (p.is_file(), p.name.lower()))
        except PermissionError:
            return ToolResult.text(f"Permission denied: {path}", is_error=True)
        lines = []
        for e in entries:
            try:
                if e.is_dir():
                    lines.append(f"[dir]  {e.name}/")
                else:
                    lines.append(f"       {e.name}  ({e.stat().st_size} bytes)")
            except OSError:
                lines.append(f"       {e.name}")
        return ToolResult.text(f"{path}\n" + ("\n".join(lines) if lines else "(empty)"))


def _find_files(roots: list[Path], pattern: str, max_results: int) -> list[str]:
    matches: list[str] = []
    for root in roots:
        for dirpath, dirnames, filenames in os.walk(root, topdown=True, onerror=lambda e: None):
            # prune noise/heavy dirs in place (speed only; not an access restriction)
            dirnames[:] = [d for d in dirnames if d.lower() not in FIND_SKIP_DIRS]
            for name in filenames:
                if fnmatch.fnmatch(name, pattern):
                    matches.append(str(Path(dirpath) / name))
            if len(matches) >= _FIND_HARD_CAP:
                break
        if len(matches) >= _FIND_HARD_CAP:
            break

    # Rank so real documents and shallower (user) paths surface first — a broad
    # pattern like *CV* can't be buried under deep package/cache files.
    def _rank(p: str) -> tuple:
        return (0 if Path(p).suffix.lower() in _DOC_EXTS else 1, p.count(os.sep), p.lower())

    matches.sort(key=_rank)
    return matches[:max_results]


class FindTool(Tool):
    name = "find"
    description = (
        "Find files by name/glob pattern anywhere on the machine. By default it searches the "
        "user's folder first and then the whole machine (all drives) if nothing is found, so it "
        "is NOT limited to any one folder. Pass 'path' to target a specific root/drive (e.g. 'D:\\\\'). "
        "Use this to LOCATE a file instead of guessing its path (e.g. pattern '*CV*.docx'). "
        "Returns absolute paths."
    )
    label = "Find"
    parameters = {
        "type": "object",
        "required": ["pattern"],
        "properties": {
            "pattern": {"type": "string", "description": "Filename glob, e.g. '*.pdf' or 'CV_*.docx'."},
            "path": {"type": "string", "description": "Root directory to search (default: workspace)."},
            "max_results": {"type": "integer", "minimum": 1, "description": "Cap on results (default 100)."},
        },
    }

    def __init__(self, config):
        self.config = config

    async def execute(self, tool_call_id, params, abort, on_update=None):
        pattern = params["pattern"]
        max_results = params.get("max_results", 100)
        scope = "machine"

        if params.get("path"):
            root = _resolve(self.config, params["path"])
            if not root.is_dir():
                return ToolResult.text(f"Search root not found: {root}", is_error=True)
            matches = await asyncio.to_thread(_find_files, [root], pattern, max_results)
            scope = str(root)
        else:
            # Stage 1: the user's folder (fast, covers Desktop/Documents/Downloads/OneDrive).
            home = Path(self.config.workspace)
            matches = await asyncio.to_thread(_find_files, [home], pattern, max_results)
            scope = str(home)
            # Stage 2: nothing there -> search the whole machine (all drives).
            if not matches:
                matches = await asyncio.to_thread(_find_files, _drive_roots(), pattern, max_results)
                scope = "whole machine"

        if not matches:
            return ToolResult.text(f"No files matching '{pattern}' (searched {scope}).")
        suffix = (
            f"\n(showing first {max_results}; refine pattern for more)"
            if len(matches) >= max_results
            else ""
        )
        return ToolResult.text(
            f"Found {len(matches)} match(es) for '{pattern}' (searched {scope}):\n"
            + "\n".join(matches)
            + suffix
        )
